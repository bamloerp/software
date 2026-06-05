from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / 'docs'

SOURCES = [
    DOCS / 'user-manual.md',
    DOCS / 'software-requirements-specification.md',
    DOCS / 'system-workflow.md',
]


HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)$')
ORDERED_RE = re.compile(r'^(\d+)\.\s+(.*)$')
UNORDERED_RE = re.compile(r'^-\s+(.*)$')
IMAGE_RE = re.compile(r'^!\[(.*?)\]\((.*?)\)$')


def bookmark_name(index: int) -> str:
    return f'toc_heading_{index}'


def apply_inline_formatting(paragraph, text: str) -> None:
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    hyperlink.set(qn('w:history'), '1')

    run = OxmlElement('w:r')
    properties = OxmlElement('w:rPr')
    style = OxmlElement('w:rStyle')
    style.set(qn('w:val'), 'Hyperlink')
    properties.append(style)
    run.append(properties)

    text_node = OxmlElement('w:t')
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_base_styles(document: Document) -> None:
    normal = document.styles['Normal']
    normal.font.name = 'Aptos'
    normal.font.size = Pt(11)

    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def collect_headings(markdown_path: Path) -> list[tuple[int, int, str]]:
    headings: list[tuple[int, int, str]] = []
    for line in markdown_path.read_text(encoding='utf-8').splitlines():
      match = HEADING_RE.match(line.strip())
      if not match:
          continue
      level = min(len(match.group(1)), 4)
      headings.append((len(headings) + 1, level, match.group(2)))
    return headings


def add_table_of_contents(document: Document, headings: list[tuple[int, int, str]]) -> None:
    document.add_heading('Table of Contents', level=1)
    for index, level, text in headings:
        if level == 1:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(max(level - 2, 0) * 0.25)
        add_internal_hyperlink(paragraph, text, bookmark_name(index))
    document.add_page_break()


def add_markdown(document: Document, markdown_path: Path) -> None:
    lines = markdown_path.read_text(encoding='utf-8').splitlines()
    heading_index = 0

    for line in lines:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph('')
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            heading_index += 1
            level = min(len(heading_match.group(1)), 4)
            heading = document.add_heading(level=level)
            apply_inline_formatting(heading, heading_match.group(2))
            add_bookmark(heading, bookmark_name(heading_index), heading_index)
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            alt_text, rel_path = image_match.groups()
            image_path = (markdown_path.parent / rel_path).resolve()
            if image_path.exists():
                if document.paragraphs and document.paragraphs[-1].text == '':
                    document.paragraphs[-1]._element.getparent().remove(document.paragraphs[-1]._element)
                document.add_picture(str(image_path), width=Inches(6.3))
                caption = document.add_paragraph()
                caption.style = 'Caption'
                caption.alignment = 1
                caption.add_run(alt_text)
            else:
                missing = document.add_paragraph()
                missing.add_run(f'[Missing image: {rel_path}]').italic = True
            continue

        ordered_match = ORDERED_RE.match(stripped)
        if ordered_match:
            paragraph = document.add_paragraph(style='List Number')
            apply_inline_formatting(paragraph, ordered_match.group(2))
            continue

        unordered_match = UNORDERED_RE.match(stripped)
        if unordered_match:
            paragraph = document.add_paragraph(style='List Bullet')
            apply_inline_formatting(paragraph, unordered_match.group(1))
            continue

        paragraph = document.add_paragraph()
        apply_inline_formatting(paragraph, stripped)


def export_docx(markdown_path: Path) -> Path:
    document = Document()
    set_base_styles(document)
    add_table_of_contents(document, collect_headings(markdown_path))
    add_markdown(document, markdown_path)
    output_path = markdown_path.with_suffix('.docx')
    try:
        document.save(output_path)
    except PermissionError:
        output_path = markdown_path.with_name(f'{markdown_path.stem}-with-toc.docx')
        document.save(output_path)
    return output_path


def main() -> None:
    exported = [export_docx(path) for path in SOURCES]
    for path in exported:
        print(path)


if __name__ == '__main__':
    main()